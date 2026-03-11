#comment
import os
from docx import Document
from datetime import datetime
from huggingface_hub import InferenceClient
import winreg
import easyocr
import PIL

client = InferenceClient(model="Qwen/Qwen2.5-72B-Instruct", token="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")

try:
    reader = easyocr.Reader(['ru', 'en'])
    print("MEIKA AI: EasyOCR успешно загружен!")
except Exception as e:
    print(f"MEIKA AI: Ошибка загрузки EasyOCR: {e}")
    reader = None

def get_safe_filename(filename):
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    return filename

def get_desktop_folder():
    try:
        key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, 
                            r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders')
        desktop_path = winreg.QueryValueEx(key, 'Desktop')[0]
        winreg.CloseKey(key)
        return desktop_path
    except:
        user_profile = os.environ.get('USERPROFILE', '')
        desktop = os.path.join(user_profile, 'Desktop')
        return desktop if os.path.exists(desktop) else os.getcwd()

def generate_summary_from_image(image_filename):
    try:
        if reader is None:
            return "Ошибка: EasyOCR не загружен"
        
        if not os.path.exists(image_filename):
            return "Ошибка: файл не найден"
        
        print("MEIKA AI: Распознаю текст с изображения с помощью EasyOCR...")
        
        results = reader.readtext(image_filename)
        
        extracted_text = ' '.join([result[1] for result in results])
        
        if not extracted_text.strip():
            return "Ошибка: не удалось распознать текст на изображении"
        
        print(f"MEIKA AI: Распознано {len(extracted_text)} символов")
        print(f"MEIKA AI: Текст: {extracted_text[:200]}...")

        prompt = f"""Создай подробный структурированный конспект на основе распознанного текста с изображения:

{extracted_text}

Используй форматирование с заголовками, подзаголовками и списками. Выдели основные идеи и ключевые моменты."""

        response = client.chat_completion(
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2000
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        return f"Ошибка при обработке изображения: {str(e)}"

def list_txt_files():
    return [f for f in os.listdir() if f.endswith('.txt')]

def list_image_files():
    image_extensions = ['.jpg', '.jpeg', '.png', '.bmp', '.gif']
    return [f for f in os.listdir() if any(f.lower().endswith(ext) for ext in image_extensions)]

def chat_with_ai(message):
    try:
        response = client.chat_completion(
            messages=[{"role": "user", "content": message}],
            max_tokens=750
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Ошибка при общении с ИИ: {e}"

def generate_summary_from_file(txt_filename):
    try:
        with open(txt_filename, 'r', encoding='utf-8') as file:
            file_content = file.read()
        
        prompt = f"""Создай подробный структурированный конспект на основе следующего текста:

{file_content}

Используй форматирование с заголовками и списками."""

        response = client.chat_completion(
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2000
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Ошибка: {e}"

def save_to_docx(content, filename):
    try:
        safe_filename = get_safe_filename(filename)
        desktop_folder = get_desktop_folder()
        full_path = os.path.join(desktop_folder, safe_filename)
        
        counter = 1
        base_name = safe_filename.replace('.docx', '')
        while os.path.exists(full_path):
            safe_filename = f"{base_name}_{counter}.docx"
            full_path = os.path.join(desktop_folder, safe_filename)
            counter += 1
        
        doc = Document()
        doc.add_heading('Конспект', 0)
        doc.add_paragraph(f"Создан: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        doc.add_paragraph(content)
        doc.save(full_path)
        return True, full_path
    except Exception as e:
        print(f"Ошибка при сохранении: {e}")
        return False, None

def choose_file_interactive(file_type="txt"):
    if file_type == "txt":
        files = list_txt_files()
        file_type_description = "текстовых"
    else:
        files = list_image_files()
        file_type_description = "изображений"
    
    if not files:
        print(f"MEIKA AI: Нет {file_type_description} файлов.")
        return None
    
    print(f"\nMEIKA AI: {file_type_description} файлы:")
    for i, filename in enumerate(files, 1):
        print(f"  {i}. {filename}")
    
    while True:
        try:
            choice = input("MEIKA AI: Введите номер или название: ").strip()
            if choice.isdigit():
                index = int(choice) - 1
                if 0 <= index < len(files):
                    return files[index]
            else:
                if file_type == "image":
                    for ext in ['.jpg', '.png', '.jpeg']:
                        if choice + ext in files:
                            return choice + ext
                else:
                    if choice + '.txt' in files:
                        return choice + '.txt'
                if choice in files:
                    return choice
            print("MEIKA AI: Файл не найден.")
        except:
            return None

def main():
    print("MEIKA AI")
    print("MEIKA AI: Теперь я могу:")
    print("- Создавать конспекты из txt файлов")
    print("- 📸 Распознавать текст с изображений")
    print("- Общаться на любые темы")
    print("\nЧем могу помочь?\n")

    while True:
        user_input = input("Вы: ").strip()

        if user_input.lower() in ['выход', 'exit']:
            print("MEIKA AI: Пока!")
            break
            
        elif user_input.lower() in ['файлы']:
            txt_files = list_txt_files()
            image_files = list_image_files()
            
            if txt_files:
                print("\nMEIKA AI: Текстовые файлы:")
                for f in txt_files:
                    print(f"  - {f}")
            
            if image_files:
                print("\nMEIKA AI: Изображения:")
                for f in image_files:
                    print(f"  - {f}")
                
        elif user_input.lower().startswith('конспект'):
            if len(user_input) > 8:
                file_name = user_input[8:].strip()
                if any(file_name.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png']):
                    file_type = "image"
                else:
                    file_type = "txt"
                    if not file_name.endswith('.txt'):
                        file_name += '.txt'
            else:
                print("MEIKA AI: Выберите тип файла:")
                print("1. Текстовый файл (.txt)")
                print("2. Изображение (.jpg, .png)")
                
                choice = input("MEIKA AI: Ваш выбор: ").strip()
                if choice == "1":
                    file_type = "txt"
                elif choice == "2":
                    file_type = "image"
                else:
                    continue
                
                file_name = choose_file_interactive(file_type)
                if not file_name:
                    continue
            
            print(f"MEIKA AI: Создаю конспект из {file_name}...")
            
            try:
                if file_type == "image":
                    summary = generate_summary_from_image(file_name)
                    output_filename = f"конспект_из_фото_{os.path.splitext(file_name)[0]}.docx"
                else:
                    summary = generate_summary_from_file(file_name)
                    output_filename = f"конспект_из_{file_name.replace('.txt', '')}.docx"

                success, saved_path = save_to_docx(summary, output_filename)
                if success:
                    print(f"MEIKA AI: ✅ Конспект успешно сохранен!")
                    print(f"MEIKA AI: 📁 Путь: {saved_path}")
                    print(f"MEIKA AI: 📊 Длина: {len(summary)} символов")
                else:
                    print("MEIKA AI: ❌ Не удалось сохранить конспект")
                    
            except Exception as e:
                print(f"MEIKA AI: Ошибка: {e}")
                
        elif user_input.lower().startswith('открой '):
            file_name = user_input.split(' ', 1)[1]
            if not file_name.endswith('.txt'):
                file_name += '.txt'
                
            try:
                with open(file_name, 'r', encoding='utf-8') as file:
                    content = file.read()
                    print(f"\nMEIKA AI: Содержимое файла '{file_name}':")
                    print("─" * 50)
                    print(content)
                    print("─" * 50)
            except:
                print(f"MEIKA AI: Файл не найден")
                
        else:
            print("MEIKA AI: Думаю...")
            response = chat_with_ai(user_input)
            print(f"MEIKA AI: {response}")

if __name__ == "__main__":
    main()